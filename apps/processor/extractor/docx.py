"""DOCX extraction module — extracts text, tables, and metadata from Word documents."""

import io
from typing import Any

from docx import Document


def extract_docx(file_bytes: bytes) -> dict[str, Any]:
    """
    Extract structured content from a DOCX file.

    Returns:
        dict with keys: text, tables, paragraphs, metadata
    """
    result: dict[str, Any] = {
        "text": "",
        "tables": [],
        "paragraphs": [],
        "metadata": {},
    }

    doc = Document(io.BytesIO(file_bytes))

    # Extract core properties
    props = doc.core_properties
    result["metadata"] = {
        "title": props.title or "",
        "author": props.author or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
    }

    # Extract paragraphs
    full_text: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            result["paragraphs"].append({
                "text": para.text,
                "style": para.style.name if para.style else "Normal",
            })

    result["text"] = "\n\n".join(full_text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        result["tables"].append({
            "index": i,
            "data": rows,
        })

    return result
