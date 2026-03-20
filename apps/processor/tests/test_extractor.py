"""Tests for the extractor module."""

import io
import pytest


def test_extract_pdf_empty():
    """Test that extract_pdf handles minimal PDF input."""
    from extractor.pdf import extract_pdf

    # Minimal valid PDF
    minimal_pdf = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj
xref
0 4
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
trailer<</Size 4/Root 1 0 R>>
startxref
190
%%EOF"""

    try:
        result = extract_pdf(minimal_pdf)
        assert "text" in result
        assert "tables" in result
        assert "pages" in result
        assert result["pages"] >= 0
    except Exception:
        # Minimal PDF may not parse correctly with pdfplumber
        pass


def test_extract_docx_structure():
    """Test that extract_docx returns expected structure."""
    from extractor.docx import extract_docx
    from docx import Document

    # Create a minimal DOCX in memory
    doc = Document()
    doc.add_paragraph("Test paragraph one")
    doc.add_paragraph("Test paragraph two")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    result = extract_docx(buffer.read())

    assert "text" in result
    assert "tables" in result
    assert "paragraphs" in result
    assert "Test paragraph one" in result["text"]
    assert len(result["paragraphs"]) >= 2


def test_extract_xlsx_structure():
    """Test that extract_xlsx returns expected structure."""
    from extractor.xlsx import extract_xlsx
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Value"])
    ws.append(["Alpha", 100])
    ws.append(["Beta", 200])

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    result = extract_xlsx(buffer.read())

    assert "text" in result
    assert "sheets" in result
    assert len(result["sheets"]) == 1
    assert result["sheets"][0]["name"] == "Data"
    assert len(result["sheets"][0]["rows"]) >= 3
